import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import os

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="SOP Formatter by S M Baqir",
    page_icon="📄",
    layout="wide"
)

# ---------------- PREMIUM CSS ----------------
st.markdown("""
<style>
.main-title{font-size:40px;font-weight:700;color:#1f4e79;}
.subtitle{color:gray;margin-bottom:25px;}
.block-container{padding-top:2rem;}
.stButton>button{background-color:#1f4e79;color:white;font-weight:600;border-radius:8px;height:45px;}
.footer{position:fixed;bottom:10px;left:0;right:0;text-align:center;color:gray;font-size:14px;}
.card{background-color:#f9fbfd;padding:20px;border-radius:10px;border:1px solid #e6eef5;}
</style>
""", unsafe_allow_html=True)

# ---------------- HEADER ----------------
st.markdown('<div class="main-title">SOP Formatter</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Professional SOP Formatting Tool — by S M Baqir</div>', unsafe_allow_html=True)

# ---------------- SIDEBAR SETTINGS ----------------
st.sidebar.title("⚙ Formatting Settings")

st.sidebar.subheader("1️⃣ First Page Settings")
first_page_font = st.sidebar.text_input("First Page Font", value="Arial")
first_page_font_size = st.sidebar.number_input("First Page Font Size", value=16)
first_page_alignment_option = st.sidebar.selectbox(
    "First Page Alignment",
    options=["Center", "Left", "Right", "Justify"],
    index=0
)

st.sidebar.subheader("2️⃣ Remaining Pages Settings")
other_page_font = st.sidebar.text_input("Body Font (2nd Page Onwards)", value="Arial")
other_page_size = st.sidebar.number_input("Font Size (2nd Page Onwards)", value=11)
line_spacing = st.sidebar.slider("Line Spacing", 1.0, 3.0, 1.5)

st.sidebar.info(
    "Bold text will be treated as headings and converted to Sentence Case from page 2 onward."
)

# ---------------- MAIN CONTENT ----------------
st.markdown('<div class="card">', unsafe_allow_html=True)
uploaded_file = st.file_uploader("Upload SOP Document (.docx)", type=["docx"])
st.markdown('</div>', unsafe_allow_html=True)

# ---------------- HELPER FUNCTIONS ----------------
def sentence_case(text):
    if not text:
        return text
    return text[0].upper() + text[1:].lower()

def get_alignment(alignment_name):
    mapping = {"Left": WD_ALIGN_PARAGRAPH.LEFT,
               "Center": WD_ALIGN_PARAGRAPH.CENTER,
               "Right": WD_ALIGN_PARAGRAPH.RIGHT,
               "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    return mapping.get(alignment_name, WD_ALIGN_PARAGRAPH.LEFT)

def set_font(run, font_name, font_size_pt):
    """Force font including bullets / numbering"""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def force_font_everywhere(doc, font_name="Arial"):
    """Apply font to all text, headers, footers, bullets"""
    for para in doc.paragraphs:
        for run in para.runs:
            size = run.font.size.pt if run.font.size else 11
            set_font(run, font_name, size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        size = run.font.size.pt if run.font.size else 11
                        set_font(run, font_name, size)

    for section in doc.sections:
        # headers
        if section.header:
            for para in section.header.paragraphs:
                for run in para.runs:
                    size = run.font.size.pt if run.font.size else 11
                    set_font(run, font_name, size)
        # footers
        if section.footer:
            for para in section.footer.paragraphs:
                for run in para.runs:
                    size = run.font.size.pt if run.font.size else 11
                    set_font(run, font_name, size)

# ---------------- FORMAT FUNCTION ----------------
def format_document(file):
    doc = Document(file)

    for i, para in enumerate(doc.paragraphs):
        # ---------------- FIRST PAGE ----------------
        if i < 10:  # approximate first page
            font = first_page_font
            size = first_page_font_size
            para.alignment = get_alignment(first_page_alignment_option)
        else:
            font = other_page_font
            size = other_page_size
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        para.paragraph_format.line_spacing = line_spacing

        for run in para.runs:
            set_font(run, font, size)
            if i >= 10 and run.bold:
                run.text = sentence_case(run.text)

    # ---------------- TABLES ----------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if i < 10:
                        font = first_page_font
                        size = first_page_font_size
                        para.alignment = get_alignment(first_page_alignment_option)
                    else:
                        font = other_page_font
                        size = other_page_size
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.line_spacing = line_spacing
                    for run in para.runs:
                        set_font(run, font, size)
                        if i >= 10 and run.bold:
                            run.text = sentence_case(run.text)

    return doc

# ---------------- PROCESS BUTTON ----------------
if uploaded_file:
    # Format button
    if st.button("✨ Format SOP Document"):
        with st.spinner("Formatting document..."):
            doc = format_document(uploaded_file)
        st.success("Document formatted successfully!")

        original_filename = uploaded_file.name
        name, ext = os.path.splitext(original_filename)
        new_filename = f"{name}_OMAC{ext}"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "⬇ Download Formatted SOP",
            data=buffer,
            file_name=new_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Force Font button
    st.markdown("---")
    st.subheader("🖋 Force Font Everywhere (Headers, Footers, Bullets, All Text)")

    force_font_name = st.text_input("Font to apply everywhere", value="Arial")

    if st.button("Apply Font to Entire Document"):
        with st.spinner("Applying font to all text, bullets, headers, and footers..."):
            doc = Document(uploaded_file)  # reload to preserve original
            force_font_everywhere(doc, force_font_name)
        st.success("Font applied everywhere successfully!")

        original_filename = uploaded_file.name
        name, ext = os.path.splitext(original_filename)
        new_filename = f"{name}_OMAC_{force_font_name}.docx"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "⬇ Download Font-Fixed Document",
            data=buffer,
            file_name=new_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ---------------- FOOTER ----------------
st.markdown('<div class="footer">OMAC Developer</div>', unsafe_allow_html=True)
